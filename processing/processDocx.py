import pandas as pd
from docx import Document
import os

csv_path = "C:/python/scripts/pdfeditor2/processing/intersection1.csv"
template_path = "C:/Users/ianm/Desktop/3J/ADA_Outputs/forms/Intersection Template.docx"
output_dir = "C:/Users/ianm/Desktop/3J/ADA_Outputs/forms"

print(f"Starting script...")

# Create output directory if it doesn't exist
os.makedirs(output_dir, exist_ok=True)

def print_table_contents(doc):
    """Debug function to print all table contents"""
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i + 1}:")
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                print(f"Row {row_idx + 1}, Cell {cell_idx + 1}: '{cell.text}'")

def replace_placeholders(doc, rows):
    print(f"\nProcessing {len(rows)} rows for placeholder replacement")
    
    for i, row in enumerate(rows.itertuples(), start=1):
        print(f"\nProcessing row {i}:")
        placeholders = {
            f'intersection{i}': str(getattr(row, 'intersection', '')),
            f'crnr_pstn{i}': str(getattr(row, 'crnr_pstn', '')),
            f'ramp_position{i}': str(getattr(row, 'ramp_position', '')),
            f'fail_count{i}': str(getattr(row, 'fail_count', '')),
            f'failing_test{i}': str(getattr(row, 'failing_test', ''))
        }
        
        print("Placeholder values for this row:")
        for key, value in placeholders.items():
            print(f"  {key}: {value}")
        
        # Process each table in the document
        for table_idx, table in enumerate(doc.tables):
            print(f"\nProcessing Table {table_idx + 1}")
            for row_idx, table_row in enumerate(table.rows):
                for cell_idx, cell in enumerate(table_row.cells):
                    original_text = cell.text
                    modified_text = original_text
                    
                    # Check each placeholder
                    for key, value in placeholders.items():
                        placeholder = f'{{{{{key}}}}}'
                        if placeholder in modified_text:
                            print(f"Found {placeholder} in Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}")
                            modified_text = modified_text.replace(placeholder, value)
                            print(f"Replaced with: {value}")
                    
                    # If we made any replacements, update the cell
                    if modified_text != original_text:
                        # Clear existing paragraphs
                        cell.paragraphs[0].clear()
                        # Add new text
                        cell.paragraphs[0].add_run(modified_text)
                        print(f"Updated cell content")
    
    return doc

try:
    # Load the CSV
    print("Attempting to read CSV...")
    df = pd.read_csv(csv_path)
    print(f"Successfully loaded CSV with {len(df)} rows")
    
    # Loop over the CSV and process each document
    for i in range(0, len(df), 4):
        print(f"\nProcessing batch starting at index {i}")
        rows = df.iloc[i:i + 4]
        
        # Load the template document
        print(f"Loading template document from: {template_path}")
        doc = Document(template_path)
        
        # Debug: Print table contents before replacement
        print("\nOriginal table contents:")
        print_table_contents(doc)
        
        # Replace placeholders
        doc = replace_placeholders(doc, rows)
        
        # Debug: Print table contents after replacement
        print("\nUpdated table contents:")
        print_table_contents(doc)
        
        try:
            intersection_name = rows.iloc[0]['intersection_name']
            print(f"Using intersection name: {intersection_name}")
            
            # Clean the filename
            safe_intersection_name = "".join(c for c in intersection_name if c.isalnum() or c in (' ', '-', '_'))
            safe_intersection_name = safe_intersection_name.replace('&', 'and')
            
            # Create the full output path
            output_path = os.path.join(output_dir, f"{safe_intersection_name}_report.docx")
            print(f"Saving document to: {output_path}")
            
            doc.save(output_path)
            print(f"Successfully saved document to: {output_path}")
            
        except Exception as e:
            print(f"Error saving document: {str(e)}")
            raise

except Exception as e:
    print(f"An error occurred: {str(e)}")
    print(f"Error type: {type(e)}")
    raise